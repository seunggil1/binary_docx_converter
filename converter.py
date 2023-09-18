import os
import base64

from docx import Document

# 바이너리 파일 <> 워드 파일간 변환 프로그램

# 라이브러리 설치 필요
# pip install python-docx

BINARYDIR = "./binary"
WORDDIR = "./docx"
UNIT = 30000000


def splitGenerator(text):
    return (text[t:t + UNIT] for t in range(0, len(text), UNIT))


def binaryToWord(fileLocation: str, fileName: str):
    global BINARYDIR
    global WORDDIR

    # load binary file
    with open(fileLocation, 'rb') as f:
        binaryData = b''.join(f.readlines())
        res = base64.b64encode(binaryData)
        res = res.decode('utf-8')

        # 파일명이랑 똑같은 directory 생성
        if not os.path.exists(f'{WORDDIR}/{fileName}'):
            os.mkdir(f'{WORDDIR}/{fileName}')

        for (idx, splitText) in enumerate(splitGenerator(res)):
            doc = Document()
            doc.add_paragraph(splitText)
            doc.save(f'{WORDDIR}/{fileName}/{fileName}_{idx}.docx')


def wordToBinary(fileLocation: str, fileList: list[str]):
    global BINARYDIR
    global WORDDIR

    with open(BINARYDIR + "/" + fileLocation[fileLocation.rfind('/') + 1:], 'wb') as g:
        for n in fileList:
            # binary to word
            doc = Document(fileLocation + '/' + n)
            data = doc.paragraphs[0].text
            data = data.encode(encoding='utf-8')
            res = base64.b64decode(data)
            g.write(res)


def converter():
    bdir = os.listdir(BINARYDIR)
    wdir = os.listdir(WORDDIR)

    for fileName in bdir:
        if fileName == ".DS_Store":  # mac에서 사용하는 메타파일 제외
            continue
        # if os.path.isfile(dir) : -> mac에서 안먹혀서 일단 제외
        print("convert binaryFile to wordFile :", fileName)
        binaryToWord(BINARYDIR + "/" + fileName, fileName)

    for fileName in wdir:
        if fileName == ".DS_Store":
            continue
        if os.path.isdir(f'{WORDDIR}/{fileName}'):
            print("convert text to visx :", fileName)
            wordToBinary(f'{WORDDIR}/{fileName}', os.listdir(f'{WORDDIR}/{fileName}'))


if __name__ == '__main__':
    converter()
